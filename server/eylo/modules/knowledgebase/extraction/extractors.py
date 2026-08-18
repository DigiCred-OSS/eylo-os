"""One extractor per format. Each takes bytes and returns text.

The output is not a faithful rendering of the document — it is the text a
retrieval index should hold. Those differ, and the differences are deliberate:
a spreadsheet becomes labelled rows rather than a grid, because "Region: EMEA,
Revenue: 41000" retrieves and a comma-separated line does not; a PDF gets page
markers, because a chunk that spans a page boundary is easier to trace back
when it says which pages it came from.

Every failure here raises `DocumentExtractionError`, which is terminal. See
`base.py` for why.
"""

from __future__ import annotations

import csv
import io
import logging

from eylo.modules.knowledgebase.extraction.base import (
    MAX_PDF_PAGES,
    MAX_SHEET_ROWS,
    MAX_TOTAL_CELLS,
    DocumentExtractionError,
    decode_text,
    guard_zip,
)

logger = logging.getLogger(__name__)


def extract_plain(raw: bytes, *, key: str) -> str:
    """Text, Markdown, reStructuredText, JSON, YAML, logs.

    Passed through unchanged. Markdown in particular is left as-is rather than
    rendered: its syntax is light enough to read, headings carry structure a
    chunker can use, and stripping it would lose exactly that.
    """
    return decode_text(raw, key=key)


# Tags that end a line of prose. Everything not listed is inline, and inline
# content must stay on one line — a sentence broken across seven lines retrieves
# as seven fragments, none of which says what the sentence said.
_BLOCK_TAGS = (
    "p",
    "div",
    "section",
    "article",
    "header",
    "footer",
    "aside",
    "main",
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "li",
    "tr",
    "br",
    "hr",
    "blockquote",
    "pre",
    "table",
    "figure",
    "figcaption",
)

# Elements whose text is never content: code the browser runs, styling, and
# fallbacks for browsers that will never see this. Minified JavaScript in
# particular is both large and superbly unhelpful — it matches tokens no human
# ever asked about.
_NON_CONTENT_TAGS = ("script", "style", "noscript", "template", "svg")


def extract_html(raw: bytes, *, key: str) -> str:
    """Visible text from an HTML document, as a reader would see it."""
    import re

    from bs4 import BeautifulSoup

    text = decode_text(raw, key=key)
    # lxml where available: faster, and far more tolerant of the broken markup
    # that real exports contain.
    try:
        soup = BeautifulSoup(text, "lxml")
    except Exception:  # noqa: BLE001 - parser availability, not document content
        soup = BeautifulSoup(text, "html.parser")

    for element in soup(list(_NON_CONTENT_TAGS)):
        element.decompose()

    title = ""
    if soup.title:
        title = soup.title.get_text(strip=True)
        soup.title.decompose()

    for element in soup.find_all(list(_BLOCK_TAGS)):
        element.append("\n")

    body = soup.get_text(separator=" ")
    lines = [re.sub(r"[^\S\n]+", " ", line).strip() for line in body.splitlines()]
    body = "\n".join(line for line in lines if line)
    return f"{title}\n\n{body}" if title else body


def extract_csv(raw: bytes, *, key: str, delimiter: str | None = None) -> str:
    """A delimited table as labelled rows.

    Each row is emitted as `header: value` pairs rather than raw fields. A bare
    CSV line loses its headers the moment it is chunked away from line one, and
    a chunk reading `EMEA, 41000, 2026-01` tells a model nothing about what
    those are. Repeating the labels costs tokens and buys the only thing that
    matters here — that a retrieved fragment is self-describing.
    """
    text = decode_text(raw, key=key)
    if delimiter is None:
        delimiter = "\t" if key.lower().endswith(".tsv") else ","

    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    try:
        first_row = next(reader, None)
        if first_row is None:
            return ""

        headers = [cell.strip() for cell in first_row]
        lines = []
        truncated = False
        for row_number, row in enumerate(reader, start=1):
            if row_number > MAX_SHEET_ROWS:
                truncated = True
                break
            pairs = [
                f"{headers[index] if index < len(headers) else f'column {index + 1}'}: {value}"
                for index, value in enumerate(row)
                if str(value).strip()
            ]
            if pairs:
                lines.append(", ".join(pairs))
    except csv.Error as error:
        raise DocumentExtractionError(
            f"{key} is not readable as CSV: {error}"
        ) from error

    if truncated:
        logger.warning(
            "Document has more than %d rows; only the first %d were indexed.",
            MAX_SHEET_ROWS,
            MAX_SHEET_ROWS,
        )
    return "\n".join(lines)


def extract_pdf(raw: bytes, *, key: str) -> str:
    """Text from a PDF, page by page."""
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(io.BytesIO(raw))
    except PdfReadError as error:
        raise DocumentExtractionError(
            f"{key} is not a readable PDF: {error}"
        ) from error

    if reader.is_encrypted:
        # An empty-string password covers the common "protected but not really"
        # case. A real password is not something this platform has.
        try:
            if not reader.decrypt(""):
                raise DocumentExtractionError(
                    f"{key} is password-protected; this platform cannot open it."
                )
        except (NotImplementedError, PdfReadError) as error:
            raise DocumentExtractionError(
                f"{key} uses an encryption this platform cannot open: {error}"
            ) from error

    pages = reader.pages
    if len(pages) > MAX_PDF_PAGES:
        logger.warning(
            "Document has %d pages; only the first %d were indexed.",
            len(pages),
            MAX_PDF_PAGES,
        )

    parts = []
    for number, page in enumerate(pages[:MAX_PDF_PAGES], start=1):
        try:
            text = page.extract_text() or ""
        except Exception as error:  # noqa: BLE001 - one bad page, not a bad file
            # Skipped rather than fatal. A single malformed page is common in
            # PDFs assembled by tooling, and losing the other 200 pages over it
            # would be the wrong trade.
            logger.warning(
                "Could not read document page=%d error_type=%s",
                number,
                type(error).__name__,
            )
            continue
        if text.strip():
            parts.append(f"[page {number}]\n{text.strip()}")

    if not parts:
        raise DocumentExtractionError(
            f"{key} contains no extractable text. It is most likely a scan, "
            "and this platform does not do OCR."
        )
    return "\n\n".join(parts)


def extract_xlsx(raw: bytes, *, key: str) -> str:
    """A modern Excel workbook as labelled rows, sheet by sheet.

    `data_only=True` reads the values Excel cached, not the formulas. A
    knowledgebase should hold what the sheet *says* — `41000` — rather than
    `=SUM(B2:B13)`, which retrieves nothing and answers nothing. The cost is
    real and worth naming: a workbook saved by a tool that never calculated it
    has no cached values, and those cells read as empty.
    """
    from openpyxl import load_workbook
    from openpyxl.utils.exceptions import InvalidFileException

    guard_zip(raw, kind="Excel")
    try:
        workbook = load_workbook(
            io.BytesIO(raw), data_only=True, read_only=True, keep_links=False
        )
    except (InvalidFileException, KeyError, ValueError) as error:
        raise DocumentExtractionError(
            f"{key} is not a readable Excel workbook: {error}"
        ) from error

    try:
        return _render_sheets(
            ((sheet.title, sheet.iter_rows(values_only=True)) for sheet in workbook),
            key=key,
        )
    finally:
        # read_only mode holds the archive open; without this the file handle
        # survives the function on some platforms.
        workbook.close()


def extract_xls(raw: bytes, *, key: str) -> str:
    """A legacy Excel workbook. Same output shape as `.xlsx`.

    Supported because "our spreadsheets" frequently means files saved a decade
    ago, and telling an operator to convert their archive by hand is not an
    answer. `xlrd` reads this format only, which is why both extractors exist.
    """
    import xlrd

    try:
        book = xlrd.open_workbook(file_contents=raw)
    except Exception as error:
        raise DocumentExtractionError(
            "Legacy Excel workbook is not readable."
        ) from error

    def sheets():
        for sheet in book.sheets():
            yield sheet.name, (sheet.row_values(index) for index in range(sheet.nrows))

    return _render_sheets(sheets(), key=key)


def extract_docx(raw: bytes, *, key: str) -> str:
    """A Word document: paragraphs and tables, in document order where possible.

    Tables are included because in practice they carry the content that matters
    — the rate card, the escalation matrix — and a parser that skipped them
    would drop the part of the document people actually ask about.
    """
    import docx
    from docx.opc.exceptions import PackageNotFoundError

    guard_zip(raw, kind="Word")
    try:
        document = docx.Document(io.BytesIO(raw))
    except (PackageNotFoundError, KeyError, ValueError) as error:
        raise DocumentExtractionError(
            f"{key} is not a readable Word document: {error}"
        ) from error

    parts = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    for table in document.tables:
        rows = table.rows
        if not rows:
            continue
        headers = [cell.text.strip() for cell in rows[0].cells]
        for row in rows[1:]:
            pairs = [
                f"{headers[index] if index < len(headers) else f'column {index + 1}'}: {cell.text.strip()}"
                for index, cell in enumerate(row.cells)
                if cell.text.strip()
            ]
            if pairs:
                parts.append(", ".join(pairs))

    if not parts:
        raise DocumentExtractionError(f"{key} contains no readable text.")
    return "\n\n".join(parts)


def extract_legacy_doc(raw: bytes, *, key: str) -> str:
    """Refuses, and says what to do instead.

    A `.doc` is an OLE compound binary, and reading it properly means either an
    external converter on the host or a fragile hand-rolled stream parser. The
    honest answer is a refusal naming the fix, not a parser that returns
    plausible-looking fragments of a document nobody can check.
    """
    raise DocumentExtractionError(
        f"{key} is a legacy Word .doc file, which this platform does not parse. "
        "Save it as .docx and re-import."
    )


def _render_sheets(sheets, *, key: str) -> str:
    """Rows from any spreadsheet, labelled with the first row's headers.

    Shared by both Excel extractors because the *output* should not depend on
    which decade the file was saved in — a caller comparing retrieval across a
    converted archive should not see two different shapes.
    """
    cells = 0
    parts: list[str] = []

    for name, rows in sheets:
        headers: list[str] = []
        lines: list[str] = []
        for index, row in enumerate(rows):
            if index > MAX_SHEET_ROWS:
                logger.warning(
                    "Sheet exceeds %d rows; the rest was not indexed.",
                    MAX_SHEET_ROWS,
                )
                break
            values = ["" if value is None else str(value).strip() for value in row]
            cells += len(values)
            if cells > MAX_TOTAL_CELLS:
                logger.warning(
                    "Document exceeds %d cells; the rest was not indexed.",
                    MAX_TOTAL_CELLS,
                )
                break
            if not any(values):
                continue
            if not headers:
                headers = values
                continue
            pairs = [
                f"{headers[position] if position < len(headers) and headers[position] else f'column {position + 1}'}: {value}"
                for position, value in enumerate(values)
                if value
            ]
            if pairs:
                lines.append(", ".join(pairs))
        if lines:
            # The sheet name is content. Workbooks routinely carry meaning in
            # tab names — "2026 Rates", "Deprecated" — and a chunk from an
            # unnamed grid loses it.
            parts.append(f"[sheet: {name}]\n" + "\n".join(lines))

    if not parts:
        raise DocumentExtractionError(f"{key} contains no readable cells.")
    return "\n\n".join(parts)
